from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .constants import *
from .models import AnswerComponent, Question
from .text_normalize import canonical_subject, compact_text, normalize_for_key, strip_question_type


def infer_subject_role(path: Path, doc: DocumentObject) -> tuple[str, str]:
    title = next((compact_text(p.text) for p in doc.paragraphs if compact_text(p.text)), path.stem)
    text = title or path.parent.name or path.stem
    role = ""
    if "教师" in text:
        role = "教师"
    elif "教研组" in text:
        role = "教研组"
    elif "学生" in text:
        role = "学生"
    elif "学校" in text:
        role = "学校"

    subject = text
    for old, new in SUBJECT_ALIASES.items():
        subject = subject.replace(old, new)
    for token in ("教师", "教研组", "学生", "学校", "（", "）", "(", ")"):
        subject = subject.replace(token, "")
    subject = subject.strip()
    if subject.endswith("学科"):
        subject = subject[: -len("学科")].strip()
    subject = subject or path.parent.name or path.stem
    subject = canonical_subject(subject)
    if subject == "学校":
        role = role or "学校"
    return subject, role


def iter_block_items(doc: DocumentObject) -> Iterable[tuple[str, Paragraph | Table]]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, doc)


def table_to_rows(table: Table) -> list[list[str]]:
    return [[compact_text(cell.text) for cell in row.cells] for row in table.rows]


def extract_question_type(text: str) -> str:
    matches = TYPE_RE.findall(text)
    if not matches:
        return ""
    last = compact_text(matches[-1])
    for known in QUESTION_TYPE_ORDER:
        if known in last:
            return known
    return last


def is_question_required(text: str) -> bool:
    return bool(QUESTION_REQUIRED_RE.match(text or ""))


def parse_dependency(text: str) -> dict[str, object] | None:
    text = compact_text(text)
    match = DEPENDENCY_RE.match(text)
    if not match:
        return None
    body = match.group(1)
    q_match = DEPENDENCY_Q_RE.search(body)
    option_orders = sorted({int(value) for value in DEPENDENCY_OPTION_RE.findall(body)})
    return {
        "dependency_text": text,
        "dependency_logic": "当且仅当" if "当且仅当" in body else "",
        "depends_on_q_no": int(q_match.group(1)) if q_match else "",
        "depends_on_option_orders": "|".join(str(value) for value in option_orders),
        "dependency_required": "需要作答" in body,
    }


def is_instruction(text: str) -> bool:
    return any(text.startswith(prefix) for prefix in NON_OPTION_PREFIXES)


def component_id(q: Question, component_type: str, order: int, label: str) -> str:
    normalized = normalize_for_key(label)[:24] or str(order)
    return f"{q.subject}_{q.q_no}_{component_type}_{order:03d}_{normalized}"


def add_paragraph_component(q: Question, text: str) -> None:
    if not text:
        return
    dependency = parse_dependency(text)
    if dependency:
        q.dependencies.append(dependency)
        return
    if q.question_type == "上传题" and text.startswith("上传限制"):
        q.answer_components.append(
            AnswerComponent(
                source_file=q.source_file,
                subject=q.subject,
                q_no=q.q_no,
                component_id=component_id(q, "upload", 1, "上传组件"),
                component_type="upload",
                component_label="上传组件",
                component_value=text,
                option_order=1,
                source_kind="paragraph",
            )
        )
        return
    if is_instruction(text):
        return
    if q.question_type in ("单选题", "多选题"):
        order = 1 + sum(1 for c in q.answer_components if c.component_type == "option")
        q.answer_components.append(
            AnswerComponent(
                source_file=q.source_file,
                subject=q.subject,
                q_no=q.q_no,
                component_id=component_id(q, "option", order, text),
                component_type="option",
                component_label=text,
                component_value=text,
                option_order=order,
                source_kind="paragraph",
            )
        )
    elif q.question_type == "填空题":
        order = 1 + sum(1 for c in q.answer_components if c.component_type == "scalar")
        q.answer_components.append(
            AnswerComponent(
                source_file=q.source_file,
                subject=q.subject,
                q_no=q.q_no,
                component_id=component_id(q, "scalar", order, text),
                component_type="scalar",
                component_label=text,
                component_value=text,
                option_order=order,
                source_kind="paragraph",
            )
        )


def add_table_components(q: Question, rows: list[list[str]]) -> None:
    if not rows:
        return
    header = rows[0]
    for col_index, raw in enumerate(header[1:], start=1):
        label = compact_text(raw)
        if not label:
            continue
        match = TABLE_CODE_RE.match(label)
        value = match.group(1) if match else ""
        clean_label = match.group(2) if match else label
        order = 1 + sum(1 for c in q.answer_components if c.component_type in ("option", "matrix_col"))
        ctype = "option" if q.question_type in ("量表题",) else "matrix_col"
        q.answer_components.append(
            AnswerComponent(
                source_file=q.source_file,
                subject=q.subject,
                q_no=q.q_no,
                component_id=component_id(q, ctype, order, clean_label),
                component_type=ctype,
                component_label=clean_label,
                component_value=value or clean_label,
                option_order=order,
                col_label=clean_label,
                source_kind="table_header",
            )
        )
    for row_index, row in enumerate(rows[1:], start=1):
        label = compact_text(row[0] if row else "")
        if not label:
            continue
        q.answer_components.append(
            AnswerComponent(
                source_file=q.source_file,
                subject=q.subject,
                q_no=q.q_no,
                component_id=component_id(q, "matrix_row", row_index, label),
                component_type="matrix_row",
                component_label=label,
                component_value=label,
                option_order=row_index,
                row_label=label,
                source_kind="table_row",
            )
        )


def extract_questions(docx_path: Path) -> list[Question]:
    doc = Document(str(docx_path))
    subject, role = infer_subject_role(docx_path, doc)
    questions: list[Question] = []
    current: Question | None = None

    def finish_current() -> None:
        if current is not None:
            questions.append(current)

    for kind, item in iter_block_items(doc):
        if kind == "paragraph":
            text = compact_text(item.text)  # type: ignore[union-attr]
            if not text:
                continue
            match = QUESTION_RE.match(text)
            matched_type = extract_question_type(match.group(2)) if match else ""
            if match and matched_type:
                finish_current()
                raw_rest = compact_text(match.group(2))
                q_type = matched_type
                q_text = strip_question_type(raw_rest)
                normalized_text = normalize_for_key(q_text)
                normalized_key = f"{role}|{q_type}|{normalized_text}"
                current = Question(
                    source_file=docx_path.name,
                    source_path=str(docx_path),
                    subject=subject,
                    role=role,
                    q_no=int(match.group(1)),
                    raw_question=text,
                    question_text=q_text,
                    question_type=q_type,
                    normalized_text=normalized_text,
                    normalized_key=normalized_key,
                    question_required=is_question_required(text),
                    block_text=[text],
                )
            elif current is not None:
                current.block_text.append(text)
                add_paragraph_component(current, text)
        elif kind == "table" and current is not None:
            rows = table_to_rows(item)  # type: ignore[arg-type]
            current.tables.append(rows)
            add_table_components(current, rows)
            flat = " | ".join(" / ".join(cell for cell in row if cell) for row in rows)
            if flat.strip():
                current.block_text.append(flat)

    finish_current()
    return questions
