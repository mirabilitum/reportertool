from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Mapping

from docx import Document


def write_report(artifact: Mapping[str, object], out_path: str | Path, template_path: str | Path | None = None) -> Path:
    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(template_path)) if template_path else Document()
    document.add_heading(text(artifact.get("title")), level=0)
    for chapter in list_items(artifact.get("chapters")):
        write_chapter(document, chapter)
    document.save(str(path))
    return path


def write_chapter(document, chapter: object) -> None:
    if not isinstance(chapter, Mapping):
        return
    document.add_heading(text(chapter.get("chapter_title")), level=1)
    for section in list_items(chapter.get("sections")):
        write_section(document, section)


def write_section(document, section: object) -> None:
    if not isinstance(section, Mapping):
        return
    document.add_heading(text(section.get("section_title")), level=2)
    for paragraph in list_items(section.get("text_paragraphs")):
        document.add_paragraph(text(paragraph))
    for result in list_items(section.get("chart_results")):
        write_chart(document, result)
    for note in list_items(section.get("notes")):
        document.add_paragraph(text(note))
    checks = list_items(section.get("quality_checks"))
    if checks:
        document.add_paragraph("质量检查")
        for check in checks:
            document.add_paragraph(format_check(check), style="List Bullet")


def write_chart(document, result: object) -> None:
    if not isinstance(result, Mapping):
        return
    caption = text(result.get("figure_caption"))
    if caption:
        document.add_paragraph(caption)
    image_path = text(result.get("image_path"))
    if image_path and Path(image_path).exists():
        document.add_picture(image_path)
        return
    encoded_png = text(result.get("encoded_png"))
    if encoded_png:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
            import base64

            temp.write(base64.b64decode(encoded_png))
            temp_path = Path(temp.name)
        try:
            document.add_picture(str(temp_path))
        finally:
            temp_path.unlink(missing_ok=True)
        return
    if result.get("svg_html", ""):
        document.add_paragraph("图表为 SVG/表格化结果，DOCX 第一版写入图题、正文、备注和质量检查；请参考 HTML 或绘图数据。")
    plot_data = text(result.get("plot_data_csv"))
    if plot_data:
        document.add_paragraph(f"绘图数据：{plot_data}")


def format_check(check: object) -> str:
    if not isinstance(check, Mapping):
        return text(check)
    parts = [text(check.get("severity")), text(check.get("check_type"))]
    prefix = " / ".join(part for part in parts if part)
    message = text(check.get("message"))
    return f"{prefix}: {message}" if prefix and message else prefix or message


def list_items(value: object) -> list:
    return value if isinstance(value, list) else []


def text(value: object) -> str:
    return "" if value is None else str(value)
