from __future__ import annotations

import json

from docx import Document

from reportertool import questionnaire_normalize


def test_questionnaire_normalize_run_writes_stage_outputs(tmp_path) -> None:
    questionnaire_dir = tmp_path / "raw"
    questionnaire_dir.mkdir()
    doc = Document()
    doc.add_paragraph("数学学科课程实施与教材使用情况表（学生）")
    doc.add_paragraph("*1. 你喜欢本学科课程吗？[单选题]")
    doc.add_paragraph("[1] 喜欢")
    doc.add_paragraph("[2] 不喜欢")
    doc.save(questionnaire_dir / "数学学生.docx")

    out_dir = tmp_path / "outputs"
    result = questionnaire_normalize.run(questionnaire_dir, out_dir)

    assert result["status"] == "passed"
    assert (out_dir / "normalized_question_mapping_long.csv").exists()
    assert (out_dir / "normalized_question_mapping_wide.csv").exists()
    assert (out_dir / "review" / "questionnaire_anomalies.csv").exists()
    status_path = out_dir / "review" / "stage_status.json"
    assert status_path.exists()
    status = json.loads(status_path.read_text(encoding="utf-8"))
    assert status["stage_name"] == "normalize-questionnaires"
    assert status["source_files"] == 1


def test_questionnaire_normalize_run_blocks_without_docx(tmp_path) -> None:
    result = questionnaire_normalize.run(tmp_path / "empty", tmp_path / "outputs")

    assert result["status"] == "blocked"
    status = json.loads((tmp_path / "outputs" / "review" / "stage_status.json").read_text(encoding="utf-8"))
    assert status["blocking_issue_count"] == 1
