from __future__ import annotations

from pathlib import Path
import unittest

from docx import Document

from questionnaire.docx_parse import infer_subject_role


class QuestionnaireSubjectParsingTest(unittest.TestCase):
    def test_student_form_title_removes_subject_suffix(self) -> None:
        doc = Document()
        doc.add_paragraph("数学学科课程实施与教材使用情况表（学生）")

        subject, role = infer_subject_role(Path("placeholder.docx"), doc)

        self.assertEqual(subject, "数学")
        self.assertEqual(role, "学生")

    def test_long_subject_student_form_keeps_canonical_subject(self) -> None:
        doc = Document()
        doc.add_paragraph("体育与健康学科课程实施与教材使用情况表（学生）")

        subject, role = infer_subject_role(Path("placeholder.docx"), doc)

        self.assertEqual(subject, "体育与健康")
        self.assertEqual(role, "学生")


if __name__ == "__main__":
    unittest.main()
